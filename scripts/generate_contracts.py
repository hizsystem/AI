#!/usr/bin/env python3
"""Generate two contract docx files for HIZ."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_chapter(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    doc.add_paragraph()


def add_article(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def add_body(doc, text):
    doc.add_paragraph(text)


def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.5
    return doc


def add_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph()


def add_parties_center(doc, gap_text, eul_text):
    for text in [gap_text, eul_text]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
    doc.add_paragraph()


def add_signature_table(doc, gap_data, eul_data):
    """gap_data and eul_data: list of [label, value] pairs"""
    table = doc.add_table(rows=1 + len(gap_data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['[ 갑 ]', '[ 을 ]']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
        set_cell_shading(cell, 'F2F2F2')

    for i, (g, e) in enumerate(zip(gap_data, eul_data)):
        table.rows[i + 1].cells[0].text = g
        table.rows[i + 1].cells[1].text = e


def add_date_line(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026년        월        일')
    run.font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph()


def add_attachment_note(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)


# ─────────────────────────────────────────────
# Contract 1: 현종용 촬영 용역 계약서
# ─────────────────────────────────────────────
def make_contract_1():
    doc = setup_doc()
    add_title(doc, '촬영 용역 계약서')
    add_parties_center(doc, '갑 : 주식회사 히즈', '을 : ________________ (현종용)')

    add_body(doc,
        '주식회사 히즈 (이하 "갑"이라 칭함)와 현종용 (이하 "을"이라 칭함)은 현대차 정몽구재단 행사 촬영 업무를 '
        '"을"이 수행함에 있어서 당사자들이 각각 성실하게 계약내용을 준수하고 상호 이익과 발전을 도모하고자 '
        '다음과 같이 촬영 용역 계약(이하 "본 계약"이라 한다)을 체결한다.')

    # Chapter 1
    add_chapter(doc, '제 1장   총칙')

    add_article(doc, '제 1조 (목적)')
    add_body(doc,
        '본 계약은 "갑"과 "을"이 상호 협력하여 현대자동차 정몽구재단(이하 "재단"이라 한다)이 주최·주관하는 '
        '각종 행사의 촬영 용역에 관하여 필요한 제반 사항을 규정함을 목적으로 한다.')

    doc.add_paragraph()
    add_article(doc, '제 2조 (상호존중)')
    add_body(doc, '① "갑"과 "을"은 본 계약에 의하여 진행되는 업무가 상호 신뢰를 기반으로 함을 인식하고 성실히 협조한다.')
    add_body(doc, '② "갑"과 "을"은 각자의 업무 영역의 고유성과 전문성을 상호 존중하여 최대한 협조한다.')

    doc.add_paragraph()
    add_article(doc, '제 3조 (용역 개요)')
    add_body(doc, '① 프로젝트명: 현대차 정몽구재단 행사 촬영 용역')
    add_body(doc, '② 계약 기간: 2026년 1월 1일 ~ 2026년 12월 31일 (총 12개월)')
    add_body(doc,
        '③ 주요 내용: 재단이 주최·주관하는 각종 행사(포럼, 콘서트, 스테이지, 정책간담회, 수료식, '
        '오리엔테이션 등)의 현장 사진 촬영 및 촬영 결과물(원본 및 보정본) 납품')
    add_body(doc, '④ 촬영 장소: 서울 및 수도권 소재 행사 현장 (구체적 장소는 행사별로 별도 안내)')
    add_body(doc,
        '⑤ 촬영 일정: "갑"이 분기별 촬영 스케줄을 사전 공유하며, 일정은 행사 사정에 따라 변동될 수 있다. '
        '일정 변경 시 "갑"은 최소 3영업일 전에 "을"에게 통보한다.')

    # Chapter 2
    add_chapter(doc, '제 2장   업무 수행 및 완료')

    add_article(doc, '제 4조 (촬영 수행)')
    add_body(doc, '① "을"은 "갑"이 통보한 일정에 따라 행사 현장에 출석하여 촬영 업무를 성실히 수행한다.')
    add_body(doc, '② "을"은 행사당 필요한 촬영 장비를 직접 준비하며, 장비에 관한 비용은 촬영비에 포함된 것으로 한다.')
    add_body(doc, '③ "을"은 촬영 완료 후 7영업일 이내에 보정된 결과물을 "갑"이 지정하는 방식으로 납품한다.')

    doc.add_paragraph()
    add_article(doc, '제 5조 (결과물 납품 및 검수)')
    add_body(doc, '① "을"이 결과물을 납품하면 "갑"은 검수할 권한이 있으며, 보완이 필요한 경우 "을"에게 보완을 요구할 수 있다.')
    add_body(doc, '② 보완 요청은 납품일로부터 5영업일 이내에 하여야 하며, "을"은 보완 요청을 받은 날로부터 3영업일 이내에 보완 결과물을 제출한다.')
    add_body(doc, '③ 검수(재검수 포함) 통과 시를 해당 건의 업무 완료시로 본다.')

    doc.add_paragraph()
    add_article(doc, '제 6조 (저작권 등 권리의 취득)')
    add_body(doc, '① 본 계약에 따라 "을"이 촬영한 결과물의 소유권 및 저작재산권은 납품 및 대금 지급 완료 시 "갑"에게 귀속한다.')
    add_body(doc,
        '② "을"은 촬영 결과물을 "을"의 포트폴리오 목적으로 사용할 수 있다. 다만 재단의 브랜드 이미지에 '
        '부정적 영향을 미칠 수 있는 방식의 사용은 금지한다.')
    add_body(doc, '③ "을"의 고유한 촬영 기법, 노하우, 작업 방법론은 "을"에게 귀속한다.')

    doc.add_paragraph()
    add_article(doc, '제 7조 (비밀 유지)')
    add_body(doc,
        '"을"은 "갑"의 사전 서면 동의 없이 본 계약에 따른 업무 수행 중 직간접적으로 알게 된 "갑" 및 재단의 '
        '기밀사항을 자기 또는 제3자의 이익을 위하여 사용하거나 제3자에게 누설해서는 아니 된다. '
        '본 조의 의무는 본 계약의 종료(해제, 해지 포함) 후에도 지속된다.')

    # Chapter 3
    add_chapter(doc, '제 3장   계약의 효력')

    add_article(doc, '제 8조 (계약 기간)')
    add_body(doc,
        '"을"의 업무 수행 기간은 2026년 1월 1일부터 2026년 12월 31일로 한다. 계약 기간 만료 1개월 전까지 '
        '양 당사자 중 일방이 서면으로 계약 종료 의사를 통지하지 않는 한, 본 계약은 동일 조건으로 1년간 자동 연장된다.')

    doc.add_paragraph()
    add_article(doc, '제 9조 (권리·의무의 양도 금지)')
    add_body(doc, '본 계약에 근거한 일체의 권리·의무에 관하여는 상대방의 사전 서면 승낙 없이 제3자에게 양도, 위탁 또는 담보로 제공할 수 없다.')

    doc.add_paragraph()
    add_article(doc, '제 10조 (계약 내용의 변경)')
    add_body(doc, '본 계약의 일부를 개정, 보완할 필요가 있을 경우에는 상호 합의하여 서면으로 변경할 수 있다.')

    doc.add_paragraph()
    add_article(doc, '제 11조 (계약의 해지)')
    add_body(doc, '① "갑" 또는 "을"은 상대방이 다음 각 호의 사유에 해당하는 경우 서면 통지로 본 계약을 해지할 수 있다.')
    add_body(doc, '    가. 본 계약을 위반하여 상대방으로부터 시정을 통보받았으나 시정 요청일로부터 10일 이내 시정이 이루어지지 않을 경우')
    add_body(doc, '    나. 해산, 파산, 폐업, 영업정지 등')
    add_body(doc, '    다. 본 계약 관련 법령 위반이나 비밀유지의무를 위반한 경우')
    add_body(doc, '② 천재지변 등으로 계약 이행이 불가능한 경우 본 계약은 해지되며, 양 당사자는 이로 인한 손해배상을 청구하지 않기로 한다.')
    add_body(doc, '③ "갑"은 경영상의 이유로 본 계약을 지속할 수 없는 경우 1개월 전 "을"에 대한 서면 통지로 본 계약을 해지할 수 있다.')
    add_body(doc, '④ 본 계약이 해지되는 경우, 해지일까지 완료된 촬영 건에 대한 대금을 정산한다.')
    add_body(doc, '⑤ 본 계약의 해지에 귀책사유가 있는 자는 이로 인하여 상대방이 입은 손해를 배상하여야 한다.')

    # Chapter 4
    add_chapter(doc, '제 4장   대금 및 절차')

    add_article(doc, '제 12조 (촬영 대금)')
    add_body(doc, '① "갑"이 "을"에게 지급하는 촬영 대금은 행사 1회당 금 440,000원 (부가가치세 포함)으로 한다.')
    add_body(doc,
        '② 연간 총 촬영 대금은 금 30,000,000원(삼천만원)을 초과하지 않는 것을 원칙으로 하며, '
        '예산 초과가 예상되는 경우 "갑"과 "을"은 사전에 협의한다.')
    add_body(doc, '③ 촬영 대금은 분기별로 정산하며, 해당 분기에 수행한 촬영 건수를 기준으로 산정한다.')

    # Payment schedule table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['구분', '지급 시기']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(cell, 'F2F2F2')

    schedule = [
        ('1분기 (1~3월)', '4월 말일까지 지급'),
        ('2분기 (4~6월)', '7월 말일까지 지급'),
        ('3분기 (7~9월)', '10월 말일까지 지급'),
        ('4분기 (10~12월)', '익년 1월 말일까지 지급'),
    ]
    for i, (q, d) in enumerate(schedule):
        table.rows[i + 1].cells[0].text = q
        table.rows[i + 1].cells[1].text = d

    doc.add_paragraph()
    add_body(doc,
        '④ 일정 변동으로 인해 촬영이 취소되는 경우, 촬영 예정일 기준 2영업일 전까지 통보하면 대금이 발생하지 않는다. '
        '2영업일 이내 취소 시 해당 건 촬영비의 50%를 "을"에게 지급한다.')

    doc.add_paragraph()
    add_article(doc, '제 13조 (대가의 청구와 지불)')
    add_body(doc,
        '"을"은 제 12조에 따라 지급할 대금을 분기 종료 후 정산하여 "갑"에게 세금계산서로 청구하고, '
        '"갑"은 세금계산서 수령일로부터 14일 이내에 "을"이 지정하는 은행 계좌에 현금으로 지급한다.')

    # Chapter 5
    add_chapter(doc, '제 5장   손해배상 등')

    add_article(doc, '제 14조 (손해배상)')
    add_body(doc,
        '당사자 중 일방이 본 계약을 위반하거나, 자신의 책임 있는 사유로 상대방에 대하여 손해를 끼친 경우 '
        '해당 당사자는 상대방에게 발생한 손해를 배상할 책임이 있다. 다만 천재지변, 비상사태 등 불가항력적 사고로 인하여 '
        '발생한 계약 불이행이나 이행의 지체에 대하여는 그 어느 일방도 상대방에게 책임을 지지 아니한다. '
        '손해배상 책임은 해당 분기 촬영 대금 총액을 한도로 한다.')

    doc.add_paragraph()
    add_article(doc, '제 15조 (제3자의 권리 침해 등)')
    add_body(doc,
        '"을"은 본 계약을 수행하는 과정에서 제3자의 소유권, 초상권, 지적재산권 등 권리를 침해하여서는 아니 되며, '
        '제3자가 권리 침해를 이유로 "갑"에게 클레임을 제기하는 경우 "을"은 "갑"을 면책하고 "을"의 비용으로 대응하여야 한다.')

    doc.add_paragraph()
    add_article(doc, '제 16조 (비상사태, 법령상의 제한 등)')
    add_body(doc,
        '천재지변, 비상사태의 발생, 법령상의 제한 등 부득이한 사유에 의해 업무 수행에 지장이 발생한 경우에는 '
        '상호 적절하고 신속한 조치를 취한다.')

    # Chapter 6
    add_chapter(doc, '제 6장   부칙')

    add_article(doc, '제 17조 (일반사항)')
    add_body(doc,
        '본 계약에 정하지 아니한 사항은 일반 상관례에 따르고, 특별히 규정해야 할 사항이 있을 때에는 '
        '별도의 규정을 상호 합의하여 첨부할 수 있다.')

    doc.add_paragraph()
    add_article(doc, '제 18조 (분쟁해결)')
    add_body(doc, '① 본 계약에 관하여 이의가 발생한 경우에는 상호 협의 후 성의를 다해 해결하는 것으로 한다.')
    add_body(doc, '② 본 계약에 관하여 분쟁이 발생한 경우에는 서울중앙지방법원을 전속적 관할법원으로 한다.')

    doc.add_paragraph()
    add_article(doc, '제 19조 (계약서의 보관)')
    add_body(doc,
        '"갑"과 "을"은 상호 대등한 입장에서 본 계약을 체결하고 성실히 계약상의 의무를 이행하기로 한다. '
        '이 증거로 본 계약서는 2부를 작성하여 서명 날인한 후 "갑"과 "을"이 각각 1부씩 보관한다.')

    add_date_line(doc)

    add_signature_table(doc,
        ['회 사 명 : 주식회사 히즈', '사업자등록번호 : 155-86-00959',
         '주    소 : 서울 마포구 양화로73-1, 5층', '대표이사 : 권 윤 정    (인)'],
        ['상    호 : ________________', '사업자등록번호 : ________________',
         '주    소 : ________________', '대 표 자 : 현 종 용    (인)'],
    )

    add_attachment_note(doc, '*첨부: 2026년 촬영 스케줄 (별첨)')

    path = '/Users/wooseongmin/Downloads/[HIZ] 촬영용역계약서_현종용_2026.docx'
    doc.save(path)
    print(f'OK: {path}')


# ─────────────────────────────────────────────
# Contract 2: 벤처리움 SNS 디자인 용역 계약서
# ─────────────────────────────────────────────
def make_contract_2():
    doc = setup_doc()
    add_title(doc, '벤처리움(Venturium)\nSNS 카드뉴스 디자인 용역 계약서')
    add_parties_center(doc, '갑 : ________________ (MYSC)', '을 : 주식회사 히즈')

    add_body(doc,
        '________________ (이하 "갑"이라 칭함)와 주식회사 히즈 (이하 "을"이라 칭함)는 "갑"이 운영하는 '
        '벤처리움(Venturium) 인스타그램 채널의 카드뉴스 디자인 용역을 "을"이 수행함에 있어서 당사자들이 '
        '각각 성실하게 계약 내용을 준수하고 상호 이익과 발전을 도모하고자 다음과 같이 디자인 용역 계약'
        '(이하 "본 계약"이라 한다)을 체결한다.')

    # Chapter 1
    add_chapter(doc, '제 1장   총칙')

    add_article(doc, '제 1조 (목적)')
    add_body(doc,
        '본 계약은 "갑"과 "을"이 상호 협력하여 한국통신사업자연합회(KTOA)가 주관하고 "갑"이 운영하는 '
        'ICT 분야 스타트업 혁신성장 지원 플랫폼 \'벤처리움(Venturium)\'의 SNS 카드뉴스 디자인 제작에 관하여 '
        '필요한 제반 사항을 규정함을 목적으로 한다.')

    doc.add_paragraph()
    add_article(doc, '제 2조 (공동사업 인식 및 상호존중)')
    add_body(doc, '① "갑"과 "을"은 본 계약에 의하여 진행되는 업무가 공동사업임을 인식하고 신뢰와 성실로써 상호 협조한다.')
    add_body(doc, '② 본 계약의 적용 및 기타 사항에 관하여 "갑"과 "을"은 각자의 업무 영역의 고유성과 전문성을 상호 존중하여 최대한 협조한다.')

    doc.add_paragraph()
    add_article(doc, '제 3조 (프로젝트 개요)')
    add_body(doc, '① 프로젝트명: 벤처리움(Venturium) SNS 카드뉴스 디자인 용역')
    add_body(doc, '② 대상 채널: KTOA 벤처리움 인스타그램')
    add_body(doc, '③ 계약 기간: 2026년 2월 2일 ~ 2026년 12월 31일')
    add_body(doc, '④ 주요 내용:')
    add_body(doc, '    가. "갑"이 제공하는 기획안을 바탕으로 카드뉴스 디자인 제작')
    add_body(doc, '    나. A타입: 기업 소식 카드뉴스')
    add_body(doc, '    다. B타입: 사업 홍보 카드뉴스')
    add_body(doc, '    라. 채널 키비주얼에 맞춘 통일감 있는 디자인 톤앤매너 구축 및 적용')
    add_body(doc, '⑤ 과업 규모: 총 15편 내외 (1편당 5~10장 구성)')
    add_body(doc, '    가. 전체 발행 장수에 따라 편수는 유동적으로 조정 가능하며, 총 납품 장수 기준으로 정산한다.')
    add_body(doc, '    나. 편수 조정이 필요한 경우 양 당사자간 사전 협의한다.')
    add_body(doc, '⑥ 작업 프로세스:')
    add_body(doc, '    1) "갑"이 기획안(텍스트, 이미지 소스 등) 제공')
    add_body(doc, '    2) "을"이 디자인 시안 제작 및 제출 (기획안 수령 후 5영업일 이내)')
    add_body(doc, '    3) "갑" 검수 및 피드백')
    add_body(doc, '    4) "을" 수정 반영 및 최종본 납품')
    add_body(doc,
        '⑦ "을"이 본 계약 기간 동안 제작한 디자인 결과물에 대한 소유권 및 저작권 등 지식재산권은 '
        '원칙적으로 "갑"에게 있으며, "을"은 "갑"이 요청한 바에 따라 결과물을 "갑"에게 제공하여야 한다.')

    # Chapter 2
    add_chapter(doc, '제 2장   업무 수행 및 완료')

    add_article(doc, '제 4조 (자료 제공)')
    add_body(doc,
        '"갑"은 업무에 필요한 기획안, 텍스트 원고, 이미지 소스, 브랜드 가이드라인 등 제반 자료를 "을"에게 제공한다. '
        '"갑"의 자료 제공 지연으로 인한 납기 지연은 "을"의 책임으로 보지 아니한다.')

    doc.add_paragraph()
    add_article(doc, '제 5조 (비밀 유지)')
    add_body(doc,
        '"을"은 "갑"의 사전 서면 동의 없이 본 계약에 따른 업무 수행 중 직간접적으로 알게 된 "갑" 및 벤처리움 관련 '
        '기밀사항을 자기 또는 제3자의 이익을 위하여 사용하거나 제3자에게 누설해서는 아니 된다. '
        '본 조의 의무는 본 계약의 종료(해제, 해지 포함)에 영향을 받지 아니하고 지속되는 것으로 한다.')

    doc.add_paragraph()
    add_article(doc, '제 6조 (사전 동의)')
    add_body(doc,
        '"을"이 제작하는 모든 디자인 결과물은 게시 전에 "갑"의 사전 서면 동의(이메일 포함)를 얻어야 한다. '
        '승인 요청 후 3영업일 이내 회신이 없는 경우 승인된 것으로 간주한다.')

    doc.add_paragraph()
    add_article(doc, '제 7조 (용역의 납품 및 검수)')
    add_body(doc,
        '① "을"이 디자인 결과물을 완료하여 납품 시 "갑"은 검수할 권한이 있으며, "갑"은 검수 후 보완이 필요하다고 '
        '판단하는 경우 "을"에게 보완을 요구할 수 있다.')
    add_body(doc,
        '② 수정 요청은 편당 2회로 한정하며, 계약 목적 범위를 벗어난 수정(콘셉트 변경, 추가 시안 등)은 '
        '추가 비용으로 별도 협의한다.')
    add_body(doc, '③ 검수 요청 후 5영업일 내 "갑"의 의견이 없으면 승인된 것으로 본다.')

    doc.add_paragraph()
    add_article(doc, '제 8조 (저작권 등 권리의 취득)')
    add_body(doc, '① 본 계약상 업무 수행 결과물의 소유권 및 저작권 등 지식재산권은 대금 지급 완료 시 전적으로 "갑"이 보유한다.')
    add_body(doc,
        '② "을"은 본 업무 수행 목적 이외의 용도로 본 계약 결과물을 사용하여서는 아니 된다. '
        '다만 "을"의 업무 수행을 알리는 홍보용 포트폴리오로 사용 가능하다.')
    add_body(doc, '③ "을"의 고유 노하우, 디자인 방법론, 포맷, 템플릿 등은 "을"에게 귀속한다.')

    # Chapter 3
    add_chapter(doc, '제 3장   계약의 효력')

    add_article(doc, '제 9조 (계약 기간)')
    add_body(doc, '"을"의 업무 수행 기간은 2026년 2월 2일부터 2026년 12월 31일로 한다. 업무 수행 일정은 상호 협의 하에 조정할 수 있다.')

    doc.add_paragraph()
    add_article(doc, '제 10조 (권리·의무의 양도 금지)')
    add_body(doc,
        '본 계약에 근거한 일체의 권리·의무에 관하여 전문 업체 활용은 사전 통보로 갈음한다. '
        '하도급 계약을 진행하는 경우에도 "갑"과 관련한 개인정보는 다루지 않는 것으로 한다.')

    doc.add_paragraph()
    add_article(doc, '제 11조 (계약 내용의 변경)')
    add_body(doc, '본 계약의 일부를 개정, 보완할 필요가 있을 경우에는 상호 합의하여 서면으로 변경할 수 있다.')

    doc.add_paragraph()
    add_article(doc, '제 12조 (변경의 통지)')
    add_body(doc,
        '"갑"과 "을"은 계약 당시 알고 있는 상호, 대표자, 소재지, 업종 및 기타 주요 사항이 변경되면 '
        '지체 없이 그 사실을 서면으로 상대방에게 통보하여야 한다.')

    doc.add_paragraph()
    add_article(doc, '제 13조 (계약의 해지)')
    add_body(doc, '① "갑" 또는 "을"은 상대방이 다음 각 호의 사유에 해당하는 경우 서면 통지로 본 계약을 해지할 수 있다.')
    add_body(doc, '    가. 본 계약을 위반하여 상대방으로부터 시정을 통보받았으나 시정 요청일로부터 10일 이내 시정이 이루어지지 않을 경우')
    add_body(doc, '    나. 해산, 파산, 폐업, 영업정지, 영업중단, 부도, 지급거절 등')
    add_body(doc, '    다. 본 계약 관련 법령 위반이나 비밀유지의무를 위반한 경우')
    add_body(doc, '② 천재지변 등으로 계약 이행이 불가능할 경우 본 계약은 해지되며, 양 당사자는 해지로 인한 손해배상을 청구하지 않기로 한다.')
    add_body(doc, '③ 본 계약이 해지되는 경우 대금은 해지일까지 완성된 업무가 본 계약 총 업무에서 해당하는 비율만큼 정산한다.')
    add_body(doc, '④ 본 계약의 해지에 귀책사유가 있는 자는 이로 인하여 상대방이 입은 손해를 배상하여야 한다.')

    # Chapter 4
    add_chapter(doc, '제 4장   대금 및 절차')

    add_article(doc, '제 14조 (대금)')
    add_body(doc, '① 본 업무를 수행한 대가로 "갑"이 "을"에게 지급하는 대금은 총 금 2,000,000원 (이백만원, 부가가치세 별도)으로 한다.')
    add_body(doc, '② 대금은 계약 체결 후 일시불로 지급하며, 세부 지급 일정은 아래와 같다.')

    # Payment table
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['구분', '금액', '지급 시기']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(cell, 'F2F2F2')

    table.rows[1].cells[0].text = '계약금 (전액)'
    table.rows[1].cells[1].text = '2,000,000원\n(VAT 별도)'
    table.rows[1].cells[2].text = '계약 체결 후\n세금계산서 발행일 기준\n14일 이내'

    doc.add_paragraph()
    add_body(doc,
        '③ 과업 규모가 당초 계약 범위(총 15편 내외)를 초과하는 경우, 양 당사자간 서면 합의로 추가 대금을 정하고 진행한다.')

    doc.add_paragraph()
    add_article(doc, '제 15조 (대가의 청구와 지불)')
    add_body(doc,
        '"을"은 제 14조에 따라 지급할 대금을 "갑"에게 세금계산서로 청구하고, "갑"은 세금계산서에 기재된 일자를 기준으로 '
        '14일 이내에 "을"이 지정하는 은행 계좌 (신한은행 : 140-012-492414 주식회사 히즈)에 현금으로 지급한다.')

    # Chapter 5
    add_chapter(doc, '제 5장   손해배상 등')

    add_article(doc, '제 16조 (손해배상)')
    add_body(doc,
        '당사자 중 일방이 본 계약을 위반하거나, 자신의 책임 있는 사유로 상대방에 대하여 손해를 끼친 경우 '
        '해당 당사자는 상대방에게 발생한 손해를 배상할 책임이 있다. 다만 천재지변, 비상사태 등 불가항력적 사고로 인하여 '
        '발생한 계약 불이행이나 이행의 지체에 대하여는 그 어느 일방도 상대방에게 책임을 지지 아니한다. '
        '손해배상 책임은 본 계약 총액을 한도로 한다.')

    doc.add_paragraph()
    add_article(doc, '제 17조 (제3자의 권리 침해 등)')
    add_body(doc,
        '"을"은 본 계약을 수행하는 과정에서 제3자의 소유권, 지적재산권 등 권리를 침해하여서는 아니 되며, '
        '제3자가 권리 침해를 이유로 "갑"에게 손해배상 등 클레임을 제기하는 경우 "을"은 "갑"을 면책하고 '
        '"을"의 비용으로 대응하여야 한다.')

    doc.add_paragraph()
    add_article(doc, '제 18조 (비상사태, 법령상의 제한 등)')
    add_body(doc,
        '천재지변, 비상사태의 발생, 법령상의 제한 등 부득이한 사유에 의해 업무 수행에 지장이 발생한 경우에는 '
        '상호 적절하고 신속한 조치를 취한다.')

    # Chapter 6
    add_chapter(doc, '제 6장   부칙')

    add_article(doc, '제 19조 (실무 담당자)')
    add_body(doc, '본 계약의 실무 커뮤니케이션은 아래 담당자를 통해 진행한다.')

    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['구분', '담당자', '연락처']):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(cell, 'F2F2F2')

    table2.rows[1].cells[0].text = '"갑" 측'
    table2.rows[1].cells[1].text = '전우철 연구원'
    table2.rows[1].cells[2].text = '________________'
    table2.rows[2].cells[0].text = '"을" 측'
    table2.rows[2].cells[1].text = '________________'
    table2.rows[2].cells[2].text = '________________'

    doc.add_paragraph()
    add_article(doc, '제 20조 (일반사항)')
    add_body(doc,
        '본 계약에 정하지 아니한 사항은 일반 상관례에 따르고, 특별히 규정해야 할 사항이 있을 때에는 '
        '별도의 규정을 상호 합의하여 첨부할 수 있다.')

    doc.add_paragraph()
    add_article(doc, '제 21조 (분쟁해결)')
    add_body(doc, '① 본 계약에 관하여 이의가 발생한 경우에는 상호 협의 후 성의를 다해 해결하는 것으로 한다.')
    add_body(doc, '② 본 계약에 관하여 분쟁이 발생한 경우에는 서울중앙지방법원을 전속적 관할법원으로 한다.')

    doc.add_paragraph()
    add_article(doc, '제 22조 (계약서의 보관)')
    add_body(doc,
        '"갑"과 "을"은 상호 대등한 입장에서 본 계약을 체결하고 성실히 계약상의 의무를 이행하기로 한다. '
        '이 증거로 본 계약서는 2부를 작성하여 서명 날인한 후 "갑"과 "을"이 각각 1부씩 보관한다.')

    add_date_line(doc)

    add_signature_table(doc,
        ['회 사 명 : ________________', '사업자등록번호 : ________________',
         '주    소 : ________________', '대표이사 : ________________    (인)'],
        ['회 사 명 : 주식회사 히즈', '사업자등록번호 : 155-86-00959',
         '주    소 : 서울 마포구 양화로73-1, 5층', '대표이사 : 권 윤 정    (인)'],
    )

    add_attachment_note(doc, '*별첨: 견적서')

    path = '/Users/wooseongmin/Downloads/[HIZ] SNS디자인용역계약서_벤더리움_2026.docx'
    doc.save(path)
    print(f'OK: {path}')


if __name__ == '__main__':
    make_contract_1()
    make_contract_2()
